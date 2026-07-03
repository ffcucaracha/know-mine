from __future__ import annotations

from pathlib import Path

from docx import Document

from src.loaders.file_router import DocumentText, PageText
from src.utils.hashing import sha256_text


def _iter_table_text(document: Document) -> list[str]:
    table_parts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cell_values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            ]
            if cell_values:
                table_parts.append(" | ".join(cell_values))
    return table_parts


def load_docx_document(path: Path) -> DocumentText:
    document = Document(path)
    paragraph_parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]
    text_parts = paragraph_parts + _iter_table_text(document)
    text = "\n".join(text_parts).strip()
    title = document.core_properties.title or None
    page = PageText(page_number=None, text=text)

    return DocumentText(
        id=sha256_text(str(path.resolve())),
        path=str(path),
        filename=path.name,
        doc_type="docx",
        title=title,
        text=text,
        pages=[page],
    )


def load_docx_text(path: Path) -> str:
    return load_docx_document(path).text
